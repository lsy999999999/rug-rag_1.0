# 文件路径: /backend/wenshu/api/autofill.py

import io
import os
from typing import List
from fastapi import APIRouter, File, Form, UploadFile, HTTPException, Depends
from fastapi.responses import StreamingResponse
from docx import Document
from datetime import datetime
from pydantic import BaseModel

from ..services.session_service import session_service
from ..processors.form_filler import DocxFormFiller
from ..llms.gpt_llm import GPTCustomLLM
from ..services.agent_service import agent_service

router = APIRouter(prefix="/autofill", tags=["Document Autofill"])

class RefineRequest(BaseModel):
    session_id: str
    feedback: str

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

async def get_gpt_llm():
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY is not configured on the server.")
    return GPTCustomLLM(model="gpt-4o")

async def _read_text(file: UploadFile) -> str:
    try:
        content_bytes = await file.read()
        await file.seek(0)
        if file.filename and file.filename.endswith('.docx'):
            doc = Document(io.BytesIO(content_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        return content_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Error reading file {file.filename}: {e}")
        return ""

@router.post("/start")
async def start_autofill_session(template_file: UploadFile = File(...)):
    if not template_file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Template must be a .docx file.")
    try:
        document = Document(io.BytesIO(await template_file.read()))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read template file: {e}")

    session_id = session_service.create_session()
    session_service.update_session(session_id, "document_state", document)
    session_service.update_session(session_id, "llm_type", "gpt")
    
    return {"session_id": session_id, "message": "Autofill session started successfully."}

async def _common_refine_logic(session_id: str, instruction: str, gpt_llm: GPTCustomLLM) -> dict:
    session = session_service.get_session(session_id)
    if not session or "document_state" not in session:
        raise HTTPException(status_code=404, detail="Session not found.")

    document = session["document_state"]
    filler = DocxFormFiller(llm=gpt_llm)
    filler.load_document(document)
    
    analysis_result = await filler.analyze_for_autofill(instruction)
    
    if not analysis_result or not analysis_result.filling_instructions:
        error_detail = analysis_result.summary if analysis_result else "Could not extract any actionable instructions."
        raise HTTPException(status_code=400, detail=error_detail)

    filler.apply_instructions([instr.model_dump() for instr in analysis_result.filling_instructions])
    session_service.update_session(session_id, "document_state", document)
    session_service.add_history(session_id, {"type": "ai_refinement", "instructions": analysis_result.model_dump()})

    return {"session_id": session_id, "message": "Feedback applied.", "preview": analysis_result.model_dump()}

@router.post("/refine")
async def refine_autofill_session(request: RefineRequest, gpt_llm = Depends(get_gpt_llm)):
    return await _common_refine_logic(request.session_id, request.feedback, gpt_llm)

@router.post("/refine_from_file")
async def refine_autofill_from_file(session_id: str = Form(...), content_file: UploadFile = File(...), gpt_llm = Depends(get_gpt_llm)):
    content_text = await _read_text(content_file)
    return await _common_refine_logic(session_id, content_text, gpt_llm)

@router.post("/refine_from_kb")
async def refine_autofill_from_kb(session_id: str = Form(...), query: str = Form(...), gpt_llm = Depends(get_gpt_llm)):
    agent = agent_service.get_agent_for_query(query)
    if not agent:
        raise HTTPException(status_code=500, detail="Could not create RAG agent.")
    response = await agent.achat(query)
    context_from_kb = str(response)
    return await _common_refine_logic(session_id, context_from_kb, gpt_llm)

@router.get("/download/{session_id}")
async def download_filled_document(session_id: str):
    session = session_service.get_session(session_id)
    if not session or "document_state" not in session:
        raise HTTPException(status_code=404, detail="Session not found.")
    document = session["document_state"]
    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    session_service.update_session(session_id, "last_updated", datetime.utcnow())
    return StreamingResponse(
        doc_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=filled_document_{session_id[:8]}.docx"}
    )

def setup_autofill_routes(app):
    app.include_router(router)