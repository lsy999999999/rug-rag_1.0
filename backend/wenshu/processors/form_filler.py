# 文件路径: /backend/wenshu/processors/form_filler.py

import json
from typing import Any, Dict, List

from docx import Document
from llama_index.core.llms.llm import LLM
from pydantic import BaseModel, Field
from openai import OpenAI

# Pydantic模型定义 (与之前一致)
class FillingInstruction(BaseModel):
    table_index: int = Field(..., description="The 0-based index of the target table.")
    row: int = Field(..., description="The 0-based row index of the target cell.")
    col: int = Field(..., description="The 0-based column index of the target cell.")
    value: str = Field(..., description="The string value to fill.")
    reason: str = Field(..., description="Explanation for choosing this location.")

class AutofillAnalysis(BaseModel):
    # We add extracted_info to match the reference code's output structure
    extracted_info: Dict[str, Any] = Field(default_factory=dict)
    filling_instructions: List[FillingInstruction] = Field(default_factory=list)
    summary: str = Field(default="Analysis complete.")


class DocxFormFiller:
    def __init__(self, llm: LLM):
        if hasattr(llm, '_client') and isinstance(llm._client, OpenAI):
            self.client = llm._client
            self.model = llm.model
        else:
            raise TypeError("DocxFormFiller requires a GPTCustomLLM instance with an active OpenAI client.")
        
        self.document: Document | None = None

    def load_document(self, doc_obj: Document):
        self.document = doc_obj

    def _extract_table_structures(self) -> List[Dict]:
        if not self.document: return []
        structures = []
        for table_idx, table in enumerate(self.document.tables):
            structure = {'table_index': table_idx, 'dimensions': f"{len(table.rows)}x{len(table.columns)}", 'cells': []}
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    structure['cells'].append({'position': f"[{row_idx},{col_idx}]", 'text': cell.text.strip(), 'is_empty': not bool(cell.text.strip())})
            structures.append(structure)
        return structures

    async def _get_filling_analysis(self, instruction: str) -> AutofillAnalysis:
        table_structures = self._extract_table_structures()
        
        # --- START: THE FINAL, CORRECTED PROMPT FROM YOUR REFERENCE ---
        system_prompt = (
            "You are an intelligent form filling assistant. You analyze table structures and determine where to fill information based on user instructions.\n\n"
            "Given:\n"
            "1. Table structure(s) with cell positions and content.\n"
            "2. User instruction with information to fill.\n\n"
            "Your task:\n"
            "1. Understand what information the user wants to fill.\n"
            "2. Find the appropriate empty cells based on nearby labels/headers.\n"
            "3. Return specific filling instructions.\n\n"
            "Rules:\n"
            "- Look for labels/headers near empty cells to determine what goes where.\n"
            "- Consider both horizontal (label to the left) and vertical (label above) layouts.\n"
            "- Be smart about matching user information to table fields even if the wording is different (e.g., '名字' and '姓名').\n"
            "- The table_index, row, and col in your output MUST be 0-based integers.\n\n"
            "Return a JSON object with this exact format:\n"
            "{\n"
            '    "extracted_info": { "field1": "value1", "field2": "value2" },\n'
            '    "filling_instructions": [\n'
            '        { "table_index": 0, "row": 0, "col": 1, "value": "value to fill", "reason": "why this location" }\n'
            '    ],\n'
            '    "summary": "Your summary of the process."\n'
            "}"
        )
        
        user_prompt = f"""Tables in the document:
{json.dumps(table_structures, ensure_ascii=False, indent=2)}

User instruction: "{instruction}"

Analyze the tables and determine where to fill the information from the user instruction."""
        # --- END: THE FINAL, CORRECTED PROMPT ---
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            result_str = response.choices[0].message.content
            print(f"[LLM RAW RESPONSE]: {result_str}")
            
            result_json = json.loads(result_str)
            # Pydantic will validate if the JSON has the correct structure.
            analysis = AutofillAnalysis.model_validate(result_json)
            return analysis
            
        except Exception as e:
            print(f"Error in LLM analysis or Pydantic validation: {e}")
            return AutofillAnalysis(filling_instructions=[], summary=f"LLM analysis failed: {e}")

    async def analyze_for_autofill(self, context: str) -> AutofillAnalysis:
        return await self._get_filling_analysis(context)

    async def refine_with_feedback(self, feedback: str) -> AutofillAnalysis:
        return await self._get_filling_analysis(feedback)

    def apply_instructions(self, instructions: List[Dict[str, Any]]):
        if not self.document: return
        for instruction in instructions:
            try:
                table_idx, row, col = int(instruction["table_index"]), int(instruction["row"]), int(instruction["col"])
                value = str(instruction.get("value", ""))
                
                if table_idx < len(self.document.tables) and row < len(self.document.tables[table_idx].rows) and col < len(self.document.tables[table_idx].columns):
                    self.document.tables[table_idx].cell(row, col).text = value
                    print(f"✓ Filled '{value}' in Table {table_idx}, Cell [{row},{col}]")
                else:
                    print(f"✗ Skipping instruction due to out-of-bounds coordinates: {instruction}")
            except Exception as e:
                print(f"✗ Failed to apply instruction {instruction}: {e}")