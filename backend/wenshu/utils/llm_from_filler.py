import os
import json
from typing import Dict, List, Any
from docx import Document
from openai import OpenAI

class AdaptiveLLMFormFiller:
    """
    支持自然语言、数据库、结构化文件自动填表的 LLM 智能表单助手
    """
    def __init__(self, api_key: str, model: str = "gpt-4o"):
        self.client = OpenAI(api_key=api_key)
        self.model = model
        self.doc = None
        self.table_structures = []

    def load_document(self, doc_path: str):
        self.doc = Document(doc_path)
        self.table_structures = self.extract_all_table_structures()

    def extract_all_table_structures(self) -> List[Dict]:
        structures = []
        for table_idx, table in enumerate(self.doc.tables):
            structure = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns) if hasattr(table, 'columns') else len(table.rows[0].cells) if table.rows else 0,
                'cells': [],
                'sample_content': []
            }
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_info = {
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell_text,
                        'is_empty': not cell_text or cell_text in ['', ' ', '\n', '\t']
                    }
                    structure['cells'].append(cell_info)
                    if not cell_info['is_empty']:
                        structure['sample_content'].append(f"[{row_idx},{col_idx}]: {cell_text}")
            structures.append(structure)
        return structures

    def analyze_and_fill_tables(self, user_instruction: str) -> Dict[str, Any]:
        tables_description = []
        for idx, structure in enumerate(self.table_structures):
            table_desc = {
                'table_number': idx + 1,
                'dimensions': f"{structure['rows']}×{structure['cols']}",
                'content': []
            }
            for cell in structure['cells']:
                if not cell['is_empty']:
                    table_desc['content'].append({
                        'position': f"[{cell['row']},{cell['col']}]",
                        'text': cell['text']
                    })
            tables_description.append(table_desc)
        system_prompt = """You are an intelligent form filling assistant. You analyze table structures and determine where to fill information based on user instructions.\n\nGiven:\n1. Table structure(s) with cell positions and content\n2. User instruction with information to fill\n\nYour task:\n1. Understand what information the user wants to fill\n2. Find the appropriate empty cells based on nearby labels/headers\n3. Return specific filling instructions\n\nRules:\n- Look for labels/headers near empty cells to determine what goes where\n- Consider both horizontal (label to the left) and vertical (label above) layouts\n- If a cell contains a label with separator (: or ：), the value goes in the same cell after the separator\n- Be smart about matching user information to table fields even if the wording is different\n\nReturn a JSON with this format:\n{\n    \"extracted_info\": {\n        \"field1\": \"value1\",\n        \"field2\": \"value2\"\n    },\n    \"filling_instructions\": [\n        {\n            \"table\": 1,\n            \"row\": 0,\n            \"col\": 1,\n            \"value\": \"value to fill\",\n            \"reason\": \"why this location\"\n        }\n    ]\n}"""
        user_prompt = f"""Tables in the document:\n{json.dumps(tables_description, ensure_ascii=False, indent=2)}\n\nUser instruction: \"{user_instruction}\"\n\nAnalyze the tables and determine where to fill the information from the user instruction."""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            result = response.choices[0].message.content.strip()
            import re
            json_match = re.search(r'\{[^{}]*\{[^{}]*\}[^{}]*\}', result, re.DOTALL)
            if not json_match:
                json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                json_str = json_str.strip()
                if '```json' in result:
                    json_match = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                elif '```' in result:
                    json_match = re.search(r'```\s*(.*?)\s*```', result, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                return json.loads(json_str)
            else:
                return {"extracted_info": {}, "filling_instructions": []}
        except Exception as e:
            return {"extracted_info": {}, "filling_instructions": []}

    def fill_cell(self, table_idx: int, row: int, col: int, value: str) -> bool:
        try:
            table = self.doc.tables[table_idx]
            cell = table.rows[row].cells[col]
            current_text = cell.text.strip()
            separators = ['：', ':', '___', '__', '_']
            for sep in separators:
                if sep in current_text:
                    parts = current_text.split(sep)
                    label_part = parts[0]
                    cell.text = f"{label_part}{sep}{value}"
                    return True
            if not current_text or current_text.replace('_', '').strip() == '':
                cell.text = value
            else:
                cell.text = f"{current_text} {value}"
            return True
        except Exception:
            return False

    def process_instruction(self, instruction: str) -> List[str]:
        result = self.analyze_and_fill_tables(instruction)
        if not result or 'filling_instructions' not in result:
            return []
        filled_items = []
        for instruction in result['filling_instructions']:
            table_idx = instruction['table'] - 1
            row = instruction['row']
            col = instruction['col']
            value = instruction['value']
            reason = instruction.get('reason', '')
            if self.fill_cell(table_idx, row, col, value):
                filled_items.append(f"✓ 填写 '{value}' 到表{instruction['table']}的[{row},{col}] - {reason}")
            else:
                filled_items.append(f"✗ 填写失败: 表{instruction['table']}的[{row},{col}]")
        return filled_items

    def save_document(self, output_path: str):
        self.doc.save(output_path)

    def fill_from_dict(self, data: Dict[str, Any]) -> List[str]:
        instruction = "，".join([f"{k}{v}" for k, v in data.items()])
        return self.process_instruction(instruction) 