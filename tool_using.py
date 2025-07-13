import os
import json
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from openai import OpenAI
import sys



class AdaptiveLLMFormFiller:
    """
    A fully adaptive form filler that uses GPT-4 to understand any table structure
    and fill it based on natural language instructions.
    """
    
    def __init__(self, api_key: str, model: str = "gpt-4o"):
        """Initialize with OpenAI API key and model"""
        self.client = OpenAI(api_key=api_key)
        self.model = model
        self.doc = None
        self.table_structures = []
    
    def load_document(self, doc_path: str):
        """Load a Word document"""
        self.doc = Document(doc_path)
        self.table_structures = self.extract_all_table_structures()
        print(f"\nLoaded document: {doc_path}")
        print(f"Found {len(self.doc.tables)} tables\n")
        
        # Display table structures
        for i, structure in enumerate(self.table_structures):
            print(f"Table {i + 1}:")
            print(f"  Dimensions: {structure['rows']} × {structure['cols']}")
            print(f"  Sample content: {structure['sample_content'][:3]}...")  # Show first 3 cells
            print()
    
    def extract_all_table_structures(self) -> List[Dict]:
        """Extract structure and content from all tables"""
        structures = []
        
        for table_idx, table in enumerate(self.doc.tables):
            structure = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns) if hasattr(table, 'columns') else len(table.rows[0].cells) if table.rows else 0,
                'cells': [],
                'sample_content': []
            }
            
            # Extract all cell information
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_info = {
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell_text,
                        'is_empty': not cell_text or cell_text in ['', ' ', '\n', '\t']
                    }
                    structure['cells'].append(cell_info)
                    
                    # Add to sample content if not empty
                    if not cell_info['is_empty']:
                        structure['sample_content'].append(f"[{row_idx},{col_idx}]: {cell_text}")
            
            structures.append(structure)
        
        return structures
    
    def analyze_and_fill_tables(self, user_instruction: str) -> Dict[str, Any]:
        """Use GPT-4 to analyze tables and determine where to fill information"""
        
        # Prepare table information for GPT-4
        tables_description = []
        for idx, structure in enumerate(self.table_structures):
            table_desc = {
                'table_number': idx + 1,
                'dimensions': f"{structure['rows']}×{structure['cols']}",
                'content': []
            }
            
            # Include all non-empty cells
            for cell in structure['cells']:
                if not cell['is_empty']:
                    table_desc['content'].append({
                        'position': f"[{cell['row']},{cell['col']}]",
                        'text': cell['text']
                    })
            
            tables_description.append(table_desc)
        
        # Create the prompt for GPT-4
        system_prompt = """You are an intelligent form filling assistant. You analyze table structures and determine where to fill information based on user instructions.

Given:
1. Table structure(s) with cell positions and content
2. User instruction with information to fill

Your task:
1. Understand what information the user wants to fill
2. Find the appropriate empty cells based on nearby labels/headers
3. Return specific filling instructions

Rules:
- Look for labels/headers near empty cells to determine what goes where
- Consider both horizontal (label to the left) and vertical (label above) layouts
- If a cell contains a label with separator (: or ：), the value goes in the same cell after the separator
- Be smart about matching user information to table fields even if the wording is different

Return a JSON with this format:
{
    "extracted_info": {
        "field1": "value1",
        "field2": "value2"
    },
    "filling_instructions": [
        {
            "table": 1,
            "row": 0,
            "col": 1,
            "value": "value to fill",
            "reason": "why this location"
        }
    ]
}"""
        
        user_prompt = f"""Tables in the document:
{json.dumps(tables_description, ensure_ascii=False, indent=2)}

User instruction: "{user_instruction}"

Analyze the tables and determine where to fill the information from the user instruction."""
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            
            result = response.choices[0].message.content.strip()
            
            # Find JSON content (might be embedded in explanation text)
            import re
            
            # Try to find JSON content between curly braces
            json_match = re.search(r'\{[^{}]*\{[^{}]*\}[^{}]*\}', result, re.DOTALL)
            if not json_match:
                # Try simpler pattern
                json_match = re.search(r'\{.*\}', result, re.DOTALL)
            
            if json_match:
                json_str = json_match.group()
                
                # Clean up the JSON string
                json_str = json_str.strip()
                
                # Remove any markdown code block markers around or within the JSON
                if '```json' in result:
                    # Extract content between ```json and ```
                    json_match = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                elif '```' in result:
                    # Extract content between ``` and ```
                    json_match = re.search(r'```\s*(.*?)\s*```', result, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                
                print(f"\nExtracted JSON:\n{json_str}\n")
                
                return json.loads(json_str)
            else:
                print(f"Could not find JSON in response:\n{result}")
                return {"extracted_info": {}, "filling_instructions": []}
            
        except Exception as e:
            print(f"Error in LLM analysis: {e}")
            return {"extracted_info": {}, "filling_instructions": []}
    
    def fill_cell(self, table_idx: int, row: int, col: int, value: str) -> bool:
        """Fill a specific cell in a table"""
        try:
            table = self.doc.tables[table_idx]
            cell = table.rows[row].cells[col]
            
            current_text = cell.text.strip()
            
            # If cell contains a separator, append value after it
            separators = ['：', ':', '___', '__', '_']
            for sep in separators:
                if sep in current_text:
                    # Keep the label and add value
                    parts = current_text.split(sep)
                    label_part = parts[0]
                    cell.text = f"{label_part}{sep}{value}"
                    return True
            
            # If cell is empty or just contains spaces/underscores
            if not current_text or current_text.replace('_', '').strip() == '':
                cell.text = value
            else:
                # Append to existing content
                cell.text = f"{current_text} {value}"
            
            return True
            
        except Exception as e:
            print(f"Error filling cell [{row},{col}] in table {table_idx + 1}: {e}")
            return False
    
    def process_instruction(self, instruction: str) -> List[str]:
        """Process a natural language instruction and fill the form"""
        print(f"\n{'='*60}")
        print(f"Processing: '{instruction}'")
        print(f"{'='*60}")
        
        # Get filling instructions from GPT-4
        result = self.analyze_and_fill_tables(instruction)
        
        if not result or 'filling_instructions' not in result:
            print("Could not determine where to fill the information.")
            return []
        
        # Execute filling instructions
        filled_items = []
        for instruction in result['filling_instructions']:
            table_idx = instruction['table'] - 1  # Convert to 0-based index
            row = instruction['row']
            col = instruction['col']
            value = instruction['value']
            reason = instruction.get('reason', '')
            
            if self.fill_cell(table_idx, row, col, value):
                filled_items.append(f"✓ Filled '{value}' in Table {instruction['table']}, Cell [{row},{col}] - {reason}")
                print(f"✓ Filled '{value}' in Table {instruction['table']}, Cell [{row},{col}]")
                print(f"  Reason: {reason}")
            else:
                print(f"✗ Failed to fill Table {instruction['table']}, Cell [{row},{col}]")
        
        return filled_items
    
    def save_document(self, output_path: str):
        """Save the filled document"""
        try:
            # Ensure directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                print(f"Created directory: {output_dir}")
            
            self.doc.save(output_path)
            print(f"\n{'='*60}")
            print(f"✓ Document saved successfully!")
            print(f"✓ File location: {output_path}")
            print(f"✓ File name: {os.path.basename(output_path)}")
            print(f"{'='*60}")
            
            # Verify file was created
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"✓ File size: {file_size:,} bytes")
            else:
                print("⚠ Warning: File may not have been saved properly!")
                
        except Exception as e:
            print(f"\n❌ Error saving document: {e}")
            print(f"❌ Attempted path: {output_path}")
    
    def show_current_state(self):
        """Display current state of all tables"""
        print("\n" + "="*60)
        print("Current Table Contents:")
        print("="*60)
        
        for table_idx, table in enumerate(self.doc.tables):
            print(f"\nTable {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_content = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        cell_text = "[empty]"
                    row_content.append(cell_text)
                print(f"  Row {row_idx}: {' | '.join(row_content)}")


def run_interactive_demo(api_key: str, doc_path: str):
    """Run an interactive demo"""
    print("\n" + "="*60)
    print("Adaptive LLM Form Filler - Works with ANY Table Structure")
    print("="*60)
    
    # Initialize the form filler
    filler = AdaptiveLLMFormFiller(api_key)
    
    # Load the document
    try:
        filler.load_document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return
    
    print("\nThis system can understand any table structure!")
    print("\nExamples of what you can say:")
    print("  - 我叫张三，电话13800138000")
    print("  - Fill in John Smith for name and john@email.com for email")
    print("  - 性别填男，年龄25岁")
    print("  - Put 'Software Engineer' in the occupation field")
    print("\nCommands:")
    print("  'show' - Display current table contents")
    print("  'quit' - Save and exit")
    print("="*60)
    
    while True:
        user_input = input("\n> ").strip()
        
        if user_input.lower() == 'quit':
            break
        
        if user_input.lower() == 'show':
            filler.show_current_state()
            continue
        
        if user_input:
            filler.process_instruction(user_input)
    
    # Save the document
    output_filename = os.path.basename(doc_path).replace('.docx', '_filled.docx')
    output_path = os.path.join('/home/lsyedith/py_test', output_filename)
    filler.save_document(output_path)
    print("\nThank you for using the Adaptive Form Filler!")


def run_batch_demo(api_key: str, doc_path: str):
    """Run a batch demo with predefined instructions"""
    print("\n" + "="*60)
    print("Batch Processing Demo")
    print("="*60)
    
    filler = AdaptiveLLMFormFiller(api_key)
    filler.load_document(doc_path)
    
    # Example batch instructions
    instructions = [
        "我的名字是张三",
        "性别男，年龄28岁",
        "电话号码13912345678",
        "邮箱地址是zhangsan@example.com",
        "工作单位是科技有限公司",
        "职位是高级工程师"
    ]
    
    print(f"\nProcessing {len(instructions)} instructions...")
    
    for instruction in instructions:
        filler.process_instruction(instruction)
    
    # Show final state
    filler.show_current_state()
    
    # Save
    output_filename = os.path.basename(doc_path).replace('.docx', '_batch_filled.docx')
    output_path = os.path.join('/home/lsyedith/py_test', output_filename)
    filler.save_document(output_path)


if __name__ == "__main__":
    # API Configuration
    API_KEY = "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"
    
    # Document path
    DOC_PATH = "/home/lsyedith/py_test/empty_list.docx"
    
    # Check if batch mode
    if len(sys.argv) > 1 and sys.argv[1] == '--batch':
        run_batch_demo(API_KEY, DOC_PATH)
    else:
        run_interactive_demo(API_KEY, DOC_PATH)




# 使用方式：
# 要确保保存文件，请确保：

# 填写完表单后键入 'quit' （不是 Ctrl+C）
# 等待保存确认消息

# 填写后的文档将被命名为：

# empty_list_filled.docx（用于交互模式）
# empty_list_batch_filled.docx （对于带有 --batch 标志的批处理模式）

# 并保存在： /home/lsyedith/py_test/


#我想要你参考这段代码，结合ruc-rag这个文件夹的内容，完成我的需求：1.使用者可以传入一个空白文档（/home/lsyedith/ruc-rag/empty_list_filled.docx），平台可以利用知识库里面的文件（/home/lsyedith/ruc-rag/test_file.docx）自动填写表格。2.请你完成这个代码，然后